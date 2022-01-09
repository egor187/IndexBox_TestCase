from pathlib import Path

import pandas
from numpy import ndarray
from docx import Document


NEW_FACTOR = 6


def dict_factory(cursor, row):
    """Func for sqlite row_factory to get dicts instead tuples"""
    d = {}
    for idx, col in enumerate(cursor.description):
        d[col[0]] = row[idx]
    return d


def calc_cagr(dataframe: pandas.DataFrame) -> ndarray:
    cagr = (
        (
            dataframe[dataframe.columns[-1:]].values[0]
            / dataframe[NEW_FACTOR][2007].values[0]
        )
        ** (1 / len(dataframe[NEW_FACTOR].columns))
        - 1
    ) * 100
    return cagr.round(2)


def word_writer(dataframe: pandas.DataFrame, path: Path, cagr: float):
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    document = Document()
    table = document.add_table(dataframe.shape[1]+1, len(dataframe.columns.names)+1)
    table.allow_autofit = True
    table.autofit = True

    for row, name in enumerate(dataframe.columns.names + list(dataframe.index)):
        table.cell(0, row).text = name

    for i, column in enumerate(dataframe, 1):
        for row in range(dataframe.shape[0]):
            table.cell(i, 1).text = str(column[1])
            table.cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.cell(i, 2).text = str(dataframe[column][row])
            table.cell(i, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    factor_cell_start = table.cell(1, 0)
    factor_cell_end = table.cell(13, 0)
    factor_cell_start.merge(factor_cell_end)
    table.cell(1, 0).text = "6"
    table.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph()
    vector = "grew" if cagr > 0 else "fall"
    document.add_paragraph(
        f"Factor 6 {vector} by avg {cagr}% every year from 2007 to 2019."
    )

    document.save(path)
