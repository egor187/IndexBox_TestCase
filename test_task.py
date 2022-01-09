import sqlite3

from pathlib import Path
import pandas
from numpy import ndarray
from docx import Document
import pandas as pd
from pandas import read_sql_query

from helpers import dict_factory

pandas.set_option("display.max_columns", 20)
pandas.set_option("display.width", 500)
NEW_FACTOR = 6
con = sqlite3.connect("test.db")
con.row_factory = dict_factory
cur = con.cursor()


def calc_cagr(dataframe: pandas.DataFrame) -> ndarray:
    cagr = (
        (
            dataframe[dataframe.columns[-1:]].values[0]
            / dataframe[NEW_FACTOR][2007].values[0]
        )
        ** (1 / len(dataframe[NEW_FACTOR].columns))
        - 1
    ) * 100

    return cagr


def word_writer(dataframe: pandas.DataFrame, path: Path, cagr: float):
    document = Document()
    table = document.add_table(dataframe.shape[0], dataframe.shape[1])
    table.allow_autofit = True
    table.autofit = True
    for i, column in enumerate(dataframe):
        for row in range(dataframe.shape[0]):
            table.cell(row, i).text = str(dataframe[column][row])

    dim = "grew" if cagr > 0 else "loss"
    document.add_paragraph(
        f"Factor 6 {dim} by avg {cagr} form 2017 to 2019"
    )

    document.save(path)


def main():
    avg_query = """
        select factor, year, sum(res) as summary from testidprod group by factor, year 
        having partner is null and state is null and bs == 0 and factor == 1 or factor == 2;
    """

    cur.execute(avg_query)
    avg_output = cur.fetchall()
    con.close()

    df = pandas.DataFrame(
        {
            (elem["factor"], elem["year"]): {"World": elem["summary"]}
            for elem in avg_output
        }
    )
    df.columns.names = ["Factor", "Year"]
    for year in [
        year["year"] for year in filter(lambda x: x["factor"] == 1, avg_output)
    ]:
        df[NEW_FACTOR, year] = df[2][year] / df[1][year]
    df.to_excel("report.xlsx")
    cagr = calc_cagr(df)

    new_factor_df = df[NEW_FACTOR]
    word_writer(new_factor_df, "report.docx", float(cagr))


if __name__ == "__main__":
    main()
